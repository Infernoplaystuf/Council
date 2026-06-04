"""Manufacturing-flavored test corpus.

Lands in ~/.council/vault/mfg_eval/ alongside the analyst_eval corpus.
Built for an F500 manufacturing analyst persona:

  parts_master.xlsx     2 sheets: Parts + BOM (multi-sheet, the canonical
                                  test for the DataIndex .xlsx patch)
  work_orders.csv       300 rows, joins parts + production lines
  defects.json          80 records, each with part_no, defect_code, area
  ecn_2026_017.pdf      3-page Engineering Change Notice (real text)
  msds_PART-3001.pdf    2-page Material Safety Data Sheet
  process_recipe.docx   text-only DOCX with a fake process recipe
  sensor_stream.csv     1000 rows of timestamped temperature data
  q3_ops_review.md      narrative referencing PART-3001 and ECN-2026-017
  images/               annotated defect photos (PNG with rendered text)
  mes.sqlite            small SQLite DB mirroring work_orders + defects
"""
from __future__ import annotations

import csv
import json
import random
import sqlite3
from datetime import date, datetime, timedelta
from pathlib import Path

import openpyxl
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import LETTER
from docx import Document

random.seed(20260603)

VAULT = Path.home() / ".council" / "vault" / "mfg_eval"
VAULT.mkdir(parents=True, exist_ok=True)
(VAULT / "images").mkdir(exist_ok=True)


PARTS = [
    ("PART-3001", "Stator Lamination Stack",   "Magnetics"),
    ("PART-3002", "Rotor Shaft Assembly",      "Machined"),
    ("PART-3003", "Housing Casting",           "Casting"),
    ("PART-3004", "Bearing Pack (sealed)",     "Purchased"),
    ("PART-3005", "Encoder PCB",               "Electronics"),
    ("PART-3006", "Terminal Box Cover",        "Sheet Metal"),
]
LINES = ["LINE-A1", "LINE-A2", "LINE-B1", "LINE-B2", "LINE-C1"]
DEFECT_CODES = [
    ("DEF-001", "Dimensional out of tolerance"),
    ("DEF-002", "Surface porosity"),
    ("DEF-003", "Insulation breakdown"),
    ("DEF-004", "Bearing noise"),
    ("DEF-005", "Color/finish nonconformance"),
    ("DEF-006", "Solder joint cracking"),
]

# ─── parts_master.xlsx (Parts + BOM, multi-sheet) ────────────
wb = Workbook()
ws = wb.active
ws.title = "Parts"
ws.append(["part_no","description","category","std_cost_usd","supplier","lead_time_days"])
for pn, desc, cat in PARTS:
    cost = round(random.uniform(120, 2200), 2)
    ws.append([pn, desc, cat, cost,
               random.choice(["Hennig", "Bosch", "TDK", "InHouse", "Sandvik"]),
               random.choice([7, 14, 21, 28, 45])])

bom = wb.create_sheet("BOM")
bom.append(["parent_part","child_part","qty","operation","work_center"])
# A small BOM tree — top assembly PART-3001 made of 3002 + 3004
bom_rows = [
    ("PART-3001","PART-3002",1,"OP-010","WC-101"),
    ("PART-3001","PART-3004",2,"OP-020","WC-105"),
    ("PART-3002","PART-3003",1,"OP-030","WC-110"),
    ("PART-3005","PART-3006",1,"OP-040","WC-115"),
]
for r in bom_rows:
    bom.append(r)
wb.save(VAULT / "parts_master.xlsx")


# ─── work_orders.csv ────────────────────────────────────────
with (VAULT / "work_orders.csv").open("w", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(["wo_id","part_no","line","start_date","qty_planned","qty_completed","status","operator"])
    for i in range(1, 301):
        pn, _, _ = random.choice(PARTS)
        planned = random.randint(50, 500)
        completed = max(0, planned - random.randint(0, 60))
        d = date(2026, 1, 1) + timedelta(days=random.randint(0, 150))
        w.writerow([f"WO-{40000+i}", pn, random.choice(LINES),
                    d.isoformat(), planned, completed,
                    random.choice(["Released","In Process","Complete","On Hold"]),
                    random.choice(["Singh","Marquez","O'Brien","Chen","Hassan"])])
# Pin a few rows for the strategic part PART-3001
with (VAULT / "work_orders.csv").open("a", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(["WO-99001","PART-3001","LINE-A1","2026-05-12",240,238,"Complete","Singh"])
    w.writerow(["WO-99002","PART-3001","LINE-A1","2026-05-19",240,210,"In Process","Singh"])
    w.writerow(["WO-99003","PART-3001","LINE-A2","2026-05-26",240,240,"Complete","Marquez"])


# ─── defects.json ───────────────────────────────────────────
defects = []
for i in range(1, 81):
    pn, _, _ = random.choice(PARTS)
    code, desc = random.choice(DEFECT_CODES)
    defects.append({
        "defect_id": f"NCR-{50000+i}",
        "part_no": pn,
        "defect_code": code,
        "description": desc,
        "qty": random.randint(1, 20),
        "logged_at": (datetime(2026,2,1) + timedelta(hours=random.randint(0, 4000))).isoformat(timespec="minutes"),
        "logged_by": random.choice(["QA1","QA2","QA3"]),
        "disposition": random.choice(["Scrap","Rework","Use-As-Is","Pending"]),
    })
# Inject a high-signal record matching the strategic part and the ECN
defects.append({
    "defect_id": "NCR-99001",
    "part_no": "PART-3001",
    "defect_code": "DEF-001",
    "description": "Dimensional out of tolerance — stack height +0.12mm beyond spec",
    "qty": 8,
    "logged_at": "2026-05-15T07:42",
    "logged_by": "QA2",
    "disposition": "Rework",
    "linked_ecn": "ECN-2026-017",
    "linked_work_order": "WO-99002",
})
(VAULT / "defects.json").write_text(json.dumps(defects, indent=2), encoding="utf-8")


# ─── ecn_2026_017.pdf (real text) ───────────────────────────
ecn = VAULT / "ecn_2026_017.pdf"
c = Canvas(str(ecn), pagesize=LETTER)
pages = [
    ("ENGINEERING CHANGE NOTICE",
     [
       "ECN Number: ECN-2026-017",
       "Effective Date: 2026-06-15",
       "Originator: M. Hassan (Manufacturing Engineering)",
       "Affected Part(s): PART-3001 (Stator Lamination Stack)",
       "",
       "Change Description:",
       "Increase stack height tolerance from +/- 0.05 mm to +/- 0.10 mm",
       "to accommodate variation observed in supplier Hennig's last 6",
       "lots. Dimensional inspection at OP-010 to remain unchanged.",
       "",
       "Driver:",
       "Non-conformance NCR-99001 logged 2026-05-15 against WO-99002.",
       "Pareto of last 12 weeks shows DEF-001 as 38% of all defects on",
       "PART-3001. Root-cause analysis at lot level points to die wear",
       "rather than process variation.",
     ]),
    ("RISK ASSESSMENT",
     [
       "Field-failure risk: low (FEA shows margin to insulation stack",
       "compression remains > 2.4x at the new tolerance ceiling).",
       "Customer-spec impact: none — internal control limit only.",
       "Regulatory impact: none.",
       "Quality plan: 100% gauging on first 3 lots post-implementation,",
       "then revert to standard AQL sampling.",
     ]),
    ("APPROVALS",
     [
       "Manufacturing Engineering: M. Hassan      2026-06-03",
       "Quality:                   J. O'Brien     2026-06-04",
       "Operations:                R. Singh       2026-06-04",
       "Customer Quality Liaison:  K. Marquez     2026-06-05",
       "",
       "Next review: ECN closure scheduled 2026-09-30 after 3 production",
       "lots have been completed and inspected per the revised plan.",
     ]),
]
for title, lines in pages:
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 730, title)
    c.setFont("Helvetica", 10)
    for i, line in enumerate(lines):
        c.drawString(72, 700 - i*15, line)
    c.showPage()
c.save()


# ─── msds_PART-3001.pdf ─────────────────────────────────────
msds = VAULT / "msds_PART-3001.pdf"
c = Canvas(str(msds), pagesize=LETTER)
c.setFont("Helvetica-Bold", 14)
c.drawString(72, 730, "MATERIAL SAFETY DATA SHEET — PART-3001")
c.setFont("Helvetica", 10)
lines = [
    "Section 1 — Identification",
    "Product: Stator Lamination Stack PART-3001",
    "Manufacturer: Acme Power Systems, Indianapolis IN",
    "Emergency contact: +1-317-555-0140",
    "",
    "Section 2 — Hazards",
    "Solvent coating on laminations is flammable; flash point 41 deg C.",
    "Avoid open flame during stack press operation OP-010.",
    "",
    "Section 8 — Exposure Controls",
    "Use local exhaust ventilation at the stack press.",
    "PPE: nitrile gloves, splash goggles when handling raw stacks.",
]
for i, line in enumerate(lines):
    c.drawString(72, 700 - i*15, line)
c.showPage()
c.save()


# ─── process_recipe.docx ────────────────────────────────────
doc = Document()
doc.add_heading("Process Recipe — Stack Press OP-010 (PART-3001)", level=1)
doc.add_paragraph("Revision 4 — supersedes Rev 3 dated 2025-09-12.")
doc.add_paragraph("Press force: 145 +/- 5 kN. Hold time: 4.2 seconds.")
doc.add_paragraph("Stack count per cycle: 24 laminations.")
doc.add_paragraph("Tooling: STP-014 (Hennig die, lot acceptance check every 1500 cycles).")
doc.add_paragraph("Post-press inspection: stack height per ECN-2026-017 revised tolerance.")
doc.add_paragraph("Operator certification required: PRO-OP-010 (annual recert).")
doc.save(VAULT / "process_recipe.docx")


# ─── sensor_stream.csv (1000 rows) ──────────────────────────
with (VAULT / "sensor_stream.csv").open("w", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(["timestamp","line","sensor","temp_c","pressure_psi","rpm","status"])
    t0 = datetime(2026, 5, 1, 6, 0)
    for i in range(1000):
        ts = t0 + timedelta(minutes=i)
        w.writerow([ts.isoformat(timespec="minutes"),
                    random.choice(LINES),
                    random.choice(["TC-01","TC-02","PR-01","RPM-01"]),
                    round(random.uniform(58, 84), 1),
                    round(random.uniform(110, 145), 1),
                    random.randint(1450, 1820),
                    random.choices(["OK","WARN","ALARM"], weights=[92,6,2])[0]])


# ─── q3_ops_review.md ───────────────────────────────────────
(VAULT / "q3_ops_review.md").write_text(
    "# Q2->Q3 2026 Ops Review (extract)\n\n"
    "PART-3001 was the long-pole this quarter. NCR-99001 (8 units, "
    "out-of-tolerance stack height) triggered ECN-2026-017 — the tolerance "
    "widening from +/- 0.05 to +/- 0.10 mm. Three work orders are tracking "
    "in-policy under the new spec (WO-99001, WO-99002, WO-99003).\n\n"
    "## Watch items\n"
    "- DEF-001 still dominates the PART-3001 Pareto (38% over 12 weeks)\n"
    "- Hennig supplier die wear is the root cause — next supplier audit Q4\n"
    "- LINE-A1 OEE held at 78%, LINE-A2 dropped to 71% — bearing-noise NCRs\n",
    encoding="utf-8",
)


# ─── annotated defect images (PNG with rendered defect text) ─
DEFECT_IMG = [
    ("ncr-99001-stack-height", "NCR-99001  PART-3001  DEF-001",
     "Stack height +0.12 mm over spec"),
    ("ncr-12044-bearing-noise", "NCR-12044  PART-3004  DEF-004",
     "Bearing noise above 72 dB at 1500 RPM"),
    ("ncr-12055-solder-crack",  "NCR-12055  PART-3005  DEF-006",
     "Solder joint crack on encoder PCB pin 4"),
]
for fname, caption, body in DEFECT_IMG:
    img = Image.new("RGB", (720, 420), color=(22, 26, 38))
    d = ImageDraw.Draw(img)
    try:
        title = ImageFont.truetype("arial.ttf", 26)
        med   = ImageFont.truetype("arial.ttf", 18)
    except Exception:
        title = med = ImageFont.load_default()
    d.text((24, 24),  caption, fill=(255, 255, 255), font=title)
    d.text((24, 70),  body,    fill=(220, 230, 255), font=med)
    d.text((24, 380), "INTERNAL — manufacturing engineering",
           fill=(140, 150, 170), font=med)
    # Fake "photo" rectangle so CLIP has some content to embed
    d.rectangle((40, 110, 680, 360), outline=(160, 200, 255), width=3)
    d.line((40, 110, 680, 360), fill=(80, 100, 140), width=1)
    img.save(VAULT / "images" / f"{fname}.png")


# ─── mes.sqlite — small warehouse for SQL connector demo ────
db_path = VAULT / "mes.sqlite"
if db_path.exists():
    db_path.unlink()
con = sqlite3.connect(db_path)
cur = con.cursor()
cur.execute("""
CREATE TABLE work_orders (
    wo_id TEXT PRIMARY KEY, part_no TEXT, line TEXT, start_date TEXT,
    qty_planned INTEGER, qty_completed INTEGER, status TEXT, operator TEXT
)""")
cur.execute("""
CREATE TABLE defects (
    defect_id TEXT PRIMARY KEY, part_no TEXT, defect_code TEXT,
    description TEXT, qty INTEGER, logged_at TEXT,
    logged_by TEXT, disposition TEXT
)""")
# Seed from the same in-memory data so cross-source queries line up.
with (VAULT / "work_orders.csv").open("r", encoding="utf-8") as f:
    r = csv.reader(f); next(r)
    cur.executemany(
        "INSERT INTO work_orders VALUES (?,?,?,?,?,?,?,?)",
        list(r),
    )
for rec in defects:
    cur.execute(
        "INSERT INTO defects VALUES (?,?,?,?,?,?,?,?)",
        (rec["defect_id"], rec["part_no"], rec["defect_code"],
         rec["description"], rec["qty"], rec["logged_at"],
         rec["logged_by"], rec["disposition"]),
    )
con.commit()
con.close()


# ─── report ─────────────────────────────────────────────────
print(f"Mfg corpus written under: {VAULT}")
total = 0
for p in sorted(VAULT.rglob("*")):
    if p.is_file():
        sz = p.stat().st_size
        total += sz
        rel = p.relative_to(VAULT)
        print(f"  {sz:>10,} B   {rel}")
print(f"  ----------\n  {total:>10,} B   total")
